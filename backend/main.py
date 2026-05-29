from fastapi import FastAPI, Depends, HTTPException, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.security import OAuth2PasswordRequestForm
from fastapi.responses import StreamingResponse, HTMLResponse
from pydantic import BaseModel
from dotenv import load_dotenv
from groq import Groq
from sqlalchemy.orm import Session
from sqlalchemy import text
from database import get_db, Base, engine
from models import User, Conversation, Message, ConversationDocument
from auth import hash_password, verify_password, create_token, get_current_user
from typing import Optional
from datetime import datetime
import os, shutil, uuid, json

import pdfplumber
from docx import Document as DocxDocument
from google.oauth2 import id_token
from google.auth.transport import requests as google_requests
from duckduckgo_search import DDGS

load_dotenv()

try:
    with engine.connect() as conn:
        conn.execute(text("ALTER TABLE users ADD COLUMN IF NOT EXISTS username VARCHAR"))
        conn.commit()
except:
    pass


client = Groq(api_key=os.getenv("GROQ_API_KEY"))
embedding_model = None
chroma_client = None
GOOGLE_CLIENT_ID = os.getenv("GOOGLE_CLIENT_ID", "")

app = FastAPI()
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173", "https://*.vercel.app", "*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

class RegisterRequest(BaseModel):
    email: str
    password: str
    username: Optional[str] = None

class ChatRequest(BaseModel):
    message: str
    conversation_id: Optional[int] = None

class GoogleLoginRequest(BaseModel):
    token: str

class UpdateProfileRequest(BaseModel):
    username: str

def get_collection_name(conv_id: int) -> str:
    return f"conv_{conv_id}"

def chunk_text(text: str, chunk_size: int = 250, overlap: int = 25):
    words = text.split()
    chunks, i = [], 0
    while i < len(words):
        chunks.append(" ".join(words[i:i + chunk_size]))
        i += chunk_size - overlap
    return chunks

def extract_text(file_path: str, filename: str) -> str:
    ext = filename.lower().split(".")[-1]
    text = ""

    if ext == "pdf":
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t: text += t + "\n"

    elif ext == "docx":
        doc = DocxDocument(file_path)
        for para in doc.paragraphs:
            if para.text.strip(): text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip(): text += cell.text + " "
                text += "\n"

    elif ext in ("txt", "py", "md", "csv", "json", "yaml", "yml",
                 "xml", "html", "js", "ts", "jsx", "tsx", "css",
                 "java", "cpp", "c", "h", "rs", "go", "rb", "php"):
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()

    elif ext == "ipynb":
        import json as jsonlib
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            nb = jsonlib.load(f)
        text += f"# Jupyter Notebook: {filename}\n\n"
        cell_num = 1
        for cell in nb.get("cells", []):
            cell_type = cell.get("cell_type", "")
            source = "".join(cell.get("source", []))
            if not source.strip():
                continue
            if cell_type == "markdown":
                text += f"## Markdown Cell {cell_num}:\n{source}\n\n"
            elif cell_type == "code":
                text += f"## Code Cell {cell_num}:\n```python\n{source}\n```\n"
                for out in cell.get("outputs", []):
                    if "text" in out:
                        out_text = "".join(out["text"])
                        if out_text.strip():
                            text += f"Output:\n{out_text}\n"
                    elif "data" in out and "text/plain" in out["data"]:
                        plain = "".join(out["data"]["text/plain"])
                        if plain.strip():
                            text += f"Result: {plain}\n"
            text += "\n"
            cell_num += 1
    else:
        raise HTTPException(status_code=400, detail=f"Unsupported file type: .{ext}")

    return text

def generate_title_from_convo(first_user: str, first_ai: str) -> str:
    try:
        preview = f"User: {first_user[:150]}\nAI: {first_ai[:150]}"
        resp = client.chat.completions.create(
            model="llama-3.1-8b-instant",
            messages=[
                {"role": "system", "content": "Generate a very short 3-5 word title summarizing this conversation. Return ONLY the title. No quotes, no punctuation at end."},
                {"role": "user", "content": preview}
            ],
            max_tokens=12
        )
        return resp.choices[0].message.content.strip()
    except:
        return first_user[:35]

def web_search(query: str) -> str:
    try:
        with DDGS() as ddgs:
            results = list(ddgs.text(query, max_results=3))
            if not results:
                return ""
            web_context = f"REAL-TIME WEB SEARCH RESULTS:\n\n"
            for i, r in enumerate(results, 1):
                body = " ".join(r['body'].split()[:100])
                web_context += f"{i}. {r['title']}: {body}\n"
            return web_context
    except Exception as e:
        print(f"Web search error: {e}")
        return ""

# ============================================================
# AUTH
# ============================================================

@app.post("/register")
def register(req: RegisterRequest, db: Session = Depends(get_db)):
    if db.query(User).filter(User.email == req.email).first():
        raise HTTPException(status_code=400, detail="Email already registered")
    username = req.username or req.email.split("@")[0]
    user = User(email=req.email, hashed_password=hash_password(req.password), username=username)
    db.add(user); db.commit()
    return {"message": "Registered successfully"}

@app.post("/login")
def login(form: OAuth2PasswordRequestForm = Depends(), db: Session = Depends(get_db)):
    user = db.query(User).filter(User.email == form.username).first()
    if not user or not verify_password(form.password, user.hashed_password):
        raise HTTPException(status_code=401, detail="Invalid credentials")
    token = create_token({"sub": user.email})
    return {"access_token": token, "token_type": "bearer"}

@app.post("/auth/google")
def google_login(req: GoogleLoginRequest, db: Session = Depends(get_db)):
    try:
        idinfo = id_token.verify_oauth2_token(req.token, google_requests.Request(), GOOGLE_CLIENT_ID)
        email = idinfo["email"]
        name = idinfo.get("name", email.split("@")[0])
    except:
        raise HTTPException(status_code=401, detail="Invalid Google token")
    user = db.query(User).filter(User.email == email).first()
    if not user:
        user = User(email=email, hashed_password=hash_password(uuid.uuid4().hex), username=name)
        db.add(user); db.commit()
    token = create_token({"sub": user.email})
    return {"access_token": token, "token_type": "bearer", "email": email, "username": user.username}

@app.get("/me")
def me(current_user: User = Depends(get_current_user)):
    return {"email": current_user.email, "username": current_user.username}

@app.put("/profile")
def update_profile(req: UpdateProfileRequest, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    current_user.username = req.username
    db.commit()
    return {"message": "Profile updated", "username": req.username}

# ============================================================
# CONVERSATIONS
# ============================================================

@app.get("/conversations")
def get_conversations(db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    convs = db.query(Conversation).filter(Conversation.user_id == current_user.id)\
        .order_by(Conversation.created_at.desc()).all()
    return [{"id": c.id, "title": c.title, "created_at": c.created_at} for c in convs]

@app.get("/conversations/{conv_id}/messages")
def get_messages(conv_id: int, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    conv = db.query(Conversation).filter(
        Conversation.id == conv_id, Conversation.user_id == current_user.id).first()
    if not conv: raise HTTPException(status_code=404, detail="Not found")
    msgs = db.query(Message).filter(Message.conversation_id == conv_id).order_by(Message.created_at).all()
    docs = db.query(ConversationDocument).filter(ConversationDocument.conversation_id == conv_id).all()
    return {
        "messages": [{"role": m.role, "content": m.content} for m in msgs],
        "documents": [d.filename for d in docs]
    }

@app.delete("/conversations/{conv_id}")
def delete_conversation(conv_id: int, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    conv = db.query(Conversation).filter(
        Conversation.id == conv_id, Conversation.user_id == current_user.id).first()
    if not conv: raise HTTPException(status_code=404, detail="Not found")
    db.query(Message).filter(Message.conversation_id == conv_id).delete()
    db.query(ConversationDocument).filter(ConversationDocument.conversation_id == conv_id).delete()
    db.delete(conv); db.commit()
    try: chroma_client.delete_collection(get_collection_name(conv_id))
    except: pass
    return {"message": "Deleted"}

@app.get("/share/{conv_id}", response_class=HTMLResponse)
def share_conversation(conv_id: int, db: Session = Depends(get_db)):
    conv = db.query(Conversation).filter(Conversation.id == conv_id).first()
    if not conv: raise HTTPException(status_code=404, detail="Not found")
    msgs = db.query(Message).filter(Message.conversation_id == conv_id).order_by(Message.created_at).all()

    import html as html_lib
    bubbles = ""
    for m in msgs:
        content = html_lib.escape(m.content).replace('\n', '<br>')
        if m.role == "user":
            bubbles += f'<div class="msg-row user-row"><div class="bubble user-bubble">{content}</div><div class="avatar user-av">U</div></div>'
        else:
            bubbles += f'<div class="msg-row ai-row"><div class="avatar ai-av">EC</div><div class="bubble ai-bubble">{content}</div></div>'

    html = f"""<!DOCTYPE html>
<html lang="en"><head>
<meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1">
<title>{html_lib.escape(conv.title)} — EasyChat</title>
<style>
*{{box-sizing:border-box;margin:0;padding:0}}
body{{font-family:Inter,system-ui,sans-serif;background:#0a0a14;color:#fff;min-height:100vh}}
.header{{background:#0d0d1a;border-bottom:1px solid #1e1e30;padding:16px 24px;display:flex;align-items:center;gap:12px;position:sticky;top:0}}
.logo{{width:36px;height:36px;background:#534AB7;border-radius:10px;display:flex;align-items:center;justify-content:center;font-weight:700;font-size:13px;color:#fff;flex-shrink:0}}
.header-title{{font-size:16px;font-weight:600}}
.header-sub{{font-size:12px;color:#555;margin-top:2px}}
.chat{{max-width:760px;margin:0 auto;padding:32px 20px}}
.msg-row{{display:flex;gap:12px;margin-bottom:24px;align-items:flex-start}}
.user-row{{flex-direction:row-reverse}}
.avatar{{width:32px;height:32px;border-radius:50%;display:flex;align-items:center;justify-content:center;font-size:11px;font-weight:700;flex-shrink:0;margin-top:2px}}
.ai-av{{background:#534AB7;color:#fff}}
.user-av{{background:#2a2a3e;color:#ccc}}
.bubble{{max-width:78%;padding:14px 18px;border-radius:18px;font-size:14px;line-height:1.7;white-space:pre-wrap;word-wrap:break-word}}
.ai-bubble{{background:#111120;border:1px solid #1e1e30;border-radius:4px 18px 18px 18px}}
.user-bubble{{background:#534AB7;border-radius:18px 4px 18px 18px}}
.footer{{text-align:center;color:#333;font-size:12px;padding:32px 0;border-top:1px solid #1e1e30;margin-top:32px}}
</style></head><body>
<div class="header">
  <div class="logo">EC</div>
  <div><div class="header-title">{html_lib.escape(conv.title)}</div>
  <div class="header-sub">Shared from EasyChat · {datetime.now().strftime("%B %d, %Y")}</div></div>
</div>
<div class="chat">{bubbles}</div>
<div class="footer">Shared from EasyChat — Your AI + RAG Pipeline</div>
</body></html>"""
    return HTMLResponse(content=html)

# ============================================================
# CHAT
# ============================================================

@app.post("/chat")
async def chat(req: ChatRequest, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    if req.conversation_id:
        conv = db.query(Conversation).filter(
            Conversation.id == req.conversation_id,
            Conversation.user_id == current_user.id).first()
        if not conv: raise HTTPException(status_code=404, detail="Not found")
    else:
        conv = Conversation(user_id=current_user.id, title="New chat")
        db.add(conv); db.commit(); db.refresh(conv)

    past = db.query(Message).filter(Message.conversation_id == conv.id)\
        .order_by(Message.created_at).all()

    conv_docs = db.query(ConversationDocument).filter(
        ConversationDocument.conversation_id == conv.id).all()
    doc_names = [d.filename for d in conv_docs]

    # RAG — strictly limit tokens
    context = ""
    if conv_docs and embedding_model and chroma_client:
        try:
            col_name = get_collection_name(conv.id)
            collection = chroma_client.get_collection(col_name)
            count = collection.count()
            if count > 0:
                q_emb = embedding_model.encode([req.message]).tolist()
                n = min(3, count)
                results = collection.query(query_embeddings=q_emb, n_results=n)
                chunks = results["documents"][0] if results["documents"] else []
                full_context = "\n\n".join(chunks)
                context = " ".join(full_context.split()[:700])
                print(f"RAG: {len(chunks)} chunks, {len(context.split())} words")
        except Exception as e:
            print(f"RAG error: {e}")

    # Web search for real-time questions
    web_context = ""
    realtime_keywords = [
        "today", "now", "current", "latest", "temperature", "weather",
        "price", "news", "right now", "live", "2025", "2026",
        "stock", "score", "result", "happened", "recently",
        "this week", "this month", "who won", "trending", "forecast"
    ]
    needs_search = any(kw in req.message.lower() for kw in realtime_keywords)
    if needs_search:
        web_context = web_search(req.message)

    today = datetime.now().strftime("%A, %B %d, %Y %I:%M %p IST")
    system_msg = f"You are EasyChat AI. Today is {today}. "

    if web_context:
        system_msg += "Use the real-time web search results to answer accurately. Give specific data. "
    if doc_names:
        system_msg += f"Files in this chat: {', '.join(doc_names)}. Answer from their content when relevant. "

    messages_list = [{"role": "system", "content": system_msg}]

    # Last 4 messages only to save tokens
    for m in past[-4:]:
        messages_list.append({"role": m.role, "content": m.content})

    # Build user content — strictly control size
    user_content = req.message
    if web_context:
        user_content += f"\n\n{' '.join(web_context.split()[:350])}"
    if context:
        user_content += f"\n\nDocument content:\n{' '.join(context.split()[:600])}"

    messages_list.append({"role": "user", "content": user_content})

    conv_id = conv.id
    is_first_msg = len(past) == 0

    db.add(Message(role="user", content=req.message, conversation_id=conv_id))
    db.commit()

    current_title = conv.title

    async def stream():
        nonlocal current_title
        full_reply = ""
        yield f"data: {json.dumps({'type': 'meta', 'conversation_id': conv_id, 'title': current_title})}\n\n"

        try:
            stream_resp = client.chat.completions.create(
                model="llama-3.1-8b-instant",
                messages=messages_list,
                stream=True,
                max_tokens=800
            )
            for chunk in stream_resp:
                delta = chunk.choices[0].delta.content or ""
                if delta:
                    full_reply += delta
                    yield f"data: {json.dumps({'type': 'token', 'token': delta})}\n\n"
        except Exception as e:
            error_msg = f"Error: {str(e)}"
            full_reply = error_msg
            yield f"data: {json.dumps({'type': 'token', 'token': error_msg})}\n\n"

        db.add(Message(role="assistant", content=full_reply, conversation_id=conv_id))
        db.commit()

        # Generate smart title AFTER first reply
        if is_first_msg and full_reply and not full_reply.startswith("Error:"):
            new_title = generate_title_from_convo(req.message, full_reply)
            conv.title = new_title
            db.commit()
            current_title = new_title
            yield f"data: {json.dumps({'type': 'title_update', 'title': new_title})}\n\n"

        yield f"data: {json.dumps({'type': 'done'})}\n\n"

    return StreamingResponse(stream(), media_type="text/event-stream")

# ============================================================
# UPLOAD
# ============================================================

@app.get("/my-documents")
def my_documents(current_user: User = Depends(get_current_user)):
    return {"documents": []}

@app.post("/upload")
async def upload_document(
    file: UploadFile = File(...),
    conversation_id: Optional[int] = Form(None),
    message: Optional[str] = Form(None),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    if conversation_id:
        conv = db.query(Conversation).filter(
            Conversation.id == conversation_id,
            Conversation.user_id == current_user.id).first()
        if not conv: raise HTTPException(status_code=404, detail="Not found")
    else:
        conv = Conversation(user_id=current_user.id, title="Document chat")
        db.add(conv); db.commit(); db.refresh(conv)

    ext = file.filename.lower().split(".")[-1]
    temp_path = f"./temp_{uuid.uuid4()}.{ext}"

    with open(temp_path, "wb") as f:
        shutil.copyfileobj(file.file, f)

    try:
        text = extract_text(temp_path, file.filename)
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))
    finally:
        if os.path.exists(temp_path):
            os.remove(temp_path)

    print(f"Extracted {len(text)} chars from {file.filename}")

    if not text.strip():
        raise HTTPException(status_code=400, detail="Could not extract text")

    chunks = chunk_text(text)
    print(f"Created {len(chunks)} chunks")
    if not embedding_model:
        raise HTTPException(
            status_code=503,
            detail="Document processing temporarily disabled on deployment server"
        )

    embeddings = embedding_model.encode(chunks).tolist()

    col_name = get_collection_name(conv.id)
    try:
        collection = chroma_client.get_collection(col_name)
    except:
        collection = chroma_client.create_collection(col_name)

    ids = [str(uuid.uuid4()) for _ in chunks]
    collection.add(documents=chunks, embeddings=embeddings, ids=ids)

    doc_record = ConversationDocument(
        conversation_id=conv.id,
        filename=file.filename,
        collection_name=col_name
    )
    db.add(doc_record); db.commit()

    ai_reply = None
    if message and message.strip():
        context = " ".join(text.split()[:1000])
        try:
            resp = client.chat.completions.create(
                model="llama-3.1-8b-instant",
                messages=[
                    {"role": "system", "content": f"You are EasyChat AI. Answer based on the uploaded file '{file.filename}'."},
                    {"role": "user", "content": f"File content:\n{context}\n\nQuestion: {message}"}
                ],
                max_tokens=700
            )
            ai_reply = resp.choices[0].message.content
            db.add(Message(role="user", content=f"[File: {file.filename}] {message}", conversation_id=conv.id))
            db.add(Message(role="assistant", content=ai_reply, conversation_id=conv.id))
            conv.title = generate_title_from_convo(message, ai_reply)
            db.commit()
        except Exception as e:
            print(f"Upload+message error: {e}")

    return {
        "message": f"Uploaded {file.filename}",
        "filename": file.filename,
        "conversation_id": conv.id,
        "ai_reply": ai_reply
    }